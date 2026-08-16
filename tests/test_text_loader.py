import tempfile
import unittest
from pathlib import Path

from docx import Document
from pypdf import PdfWriter

from app.services.text_loader import (
    EmptyDocumentError,
    MalformedDocumentError,
    UnsupportedDocumentTypeError,
    clean_text,
    extract_document,
)
from tests.document_fixtures import write_text_pdf


class TextLoaderTests(unittest.TestCase):
    def setUp(self):
        self.tempdir = tempfile.TemporaryDirectory()
        self.temp_path = Path(self.tempdir.name)

    def tearDown(self):
        self.tempdir.cleanup()

    def test_clean_text_normalizes_line_endings_and_blank_lines(self):
        self.assertEqual(
            clean_text("  First  \r\n\r\n\r\nSecond\t  \r\n"),
            "First\n\nSecond",
        )

    def test_extracts_utf8_txt(self):
        path = self.temp_path / "policy.txt"
        path.write_text("Annual leave policy", encoding="utf-8")

        document = extract_document(path)

        self.assertEqual(document.text, "Annual leave policy")
        self.assertEqual(len(document.sections), 1)
        self.assertIsNone(document.sections[0].page_number)

    def test_extracts_gbk_txt(self):
        path = self.temp_path / "policy.txt"
        path.write_bytes("年假政策".encode("gbk"))

        document = extract_document(path)

        self.assertEqual(document.text, "年假政策")

    def test_extracts_page_aware_pdf(self):
        path = self.temp_path / "policy.pdf"
        write_text_pdf(path, ["Page one policy", "Page two policy"])

        document = extract_document(path)

        self.assertEqual(
            [section.page_number for section in document.sections],
            [1, 2],
        )
        self.assertIn("[Page 1]\nPage one policy", document.text)
        self.assertIn("[Page 2]\nPage two policy", document.text)

    def test_extracts_docx_paragraphs_and_tables(self):
        path = self.temp_path / "policy.docx"
        source = Document()
        source.add_paragraph("Annual leave policy")
        table = source.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Region"
        table.cell(0, 1).text = "Days"
        source.save(path)

        document = extract_document(path)

        self.assertIn("Annual leave policy", document.text)
        self.assertIn("Region | Days", document.text)
        self.assertIsNone(document.sections[0].page_number)

    def test_missing_file_raises_file_not_found(self):
        with self.assertRaises(FileNotFoundError):
            extract_document(self.temp_path / "missing.txt")

    def test_unsupported_extension_is_rejected(self):
        path = self.temp_path / "policy.exe"
        path.write_bytes(b"content")

        with self.assertRaises(UnsupportedDocumentTypeError):
            extract_document(path)

    def test_whitespace_only_txt_is_empty_document(self):
        path = self.temp_path / "empty.txt"
        path.write_text(" \n\n\t ", encoding="utf-8")

        with self.assertRaises(EmptyDocumentError):
            extract_document(path)

    def test_blank_pdf_is_empty_document(self):
        path = self.temp_path / "empty.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        with path.open("wb") as output:
            writer.write(output)

        with self.assertRaises(EmptyDocumentError):
            extract_document(path)

    def test_empty_docx_is_empty_document(self):
        path = self.temp_path / "empty.docx"
        Document().save(path)

        with self.assertRaises(EmptyDocumentError):
            extract_document(path)

    def test_malformed_txt_pdf_and_docx_are_rejected(self):
        malformed_files = {
            "bad.txt": b"\xff",
            "bad.pdf": b"not a pdf",
            "bad.docx": b"not a docx archive",
        }

        for filename, content in malformed_files.items():
            with self.subTest(filename=filename):
                path = self.temp_path / filename
                path.write_bytes(content)
                with self.assertRaises(MalformedDocumentError):
                    extract_document(path)


if __name__ == "__main__":
    unittest.main()
