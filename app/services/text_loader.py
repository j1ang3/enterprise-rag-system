import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


class TextExtractionError(RuntimeError):
    pass


class EmptyDocumentError(TextExtractionError):
    pass


class MalformedDocumentError(TextExtractionError):
    pass


class UnsupportedDocumentTypeError(TextExtractionError):
    pass


@dataclass(frozen=True)
class ExtractedSection:
    text: str
    page_number: int | None = None


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    sections: tuple[ExtractedSection, ...]


def clean_text(text: str) -> str:
    """Normalize line endings and excessive blank lines without rewriting content."""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = "\n".join(line.rstrip() for line in normalized.split("\n"))
    return re.sub(r"\n{3,}", "\n\n", normalized).strip()


def _build_extracted_document(sections: Iterable[ExtractedSection]) -> ExtractedDocument:
    cleaned_sections = tuple(
        ExtractedSection(text=cleaned, page_number=section.page_number)
        for section in sections
        if (cleaned := clean_text(section.text))
    )
    if not cleaned_sections:
        raise EmptyDocumentError("Document does not contain extractable text.")

    if any(section.page_number is not None for section in cleaned_sections):
        rendered_sections = [
            f"[Page {section.page_number}]\n{section.text}"
            if section.page_number is not None
            else section.text
            for section in cleaned_sections
        ]
    else:
        rendered_sections = [section.text for section in cleaned_sections]

    return ExtractedDocument(
        text="\n\n".join(rendered_sections),
        sections=cleaned_sections,
    )


def read_text_file(file_path: Path) -> str:
    """
    Read text content from a .txt or .md file.

    Args:
        file_path: Path to the uploaded file.

    Returns:
        Extracted text content.
    """
    file_bytes = file_path.read_bytes()
    for encoding in ("utf-8", "gbk"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise MalformedDocumentError("Text file is not valid UTF-8 or GBK.")


def _read_pdf_sections(file_path: Path) -> tuple[ExtractedSection, ...]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise TextExtractionError(
            "PDF support requires pypdf. Install dependencies from requirements.txt."
        ) from exc

    try:
        reader = PdfReader(str(file_path))
        return tuple(
            ExtractedSection(text=page.extract_text() or "", page_number=page_number)
            for page_number, page in enumerate(reader.pages, start=1)
        )
    except Exception as exc:
        raise MalformedDocumentError("PDF document could not be parsed.") from exc


def read_pdf_file(file_path: Path) -> str:
    """
    Extract text from a PDF file, page by page.
    """
    return _build_extracted_document(_read_pdf_sections(file_path)).text


def _read_docx_sections(file_path: Path) -> tuple[ExtractedSection, ...]:
    try:
        from docx import Document
    except ImportError as exc:
        raise TextExtractionError(
            "DOCX support requires python-docx. Install dependencies from requirements.txt."
        ) from exc

    try:
        document = Document(str(file_path))
        parts = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    except Exception as exc:
        raise MalformedDocumentError("DOCX document could not be parsed.") from exc

    return (ExtractedSection(text="\n\n".join(parts)),)


def read_docx_file(file_path: Path) -> str:
    """
    Extract paragraph and table text from a DOCX file.
    """
    return _build_extracted_document(_read_docx_sections(file_path)).text


def extract_document(file_path: Path) -> ExtractedDocument:
    """Extract cleaned text plus source-location sections from a supported document."""
    if not file_path.is_file():
        raise FileNotFoundError(f"Document file not found: {file_path}")

    suffix = file_path.suffix.lower()
    if suffix in {".txt", ".md"}:
        sections = (ExtractedSection(text=read_text_file(file_path)),)
    elif suffix == ".pdf":
        sections = _read_pdf_sections(file_path)
    elif suffix == ".docx":
        sections = _read_docx_sections(file_path)
    else:
        raise UnsupportedDocumentTypeError(f"Unsupported file type: {suffix}")

    return _build_extracted_document(sections)


def extract_text(file_path: Path) -> str:
    """
    Extract text from uploaded document.

    Currently supported:
    - .txt
    - .md
    - .pdf
    - .docx
    """
    return extract_document(file_path).text


def make_preview(text: str, limit: int = 1000) -> str:
    """
    Return a shortened preview of the extracted text.
    """
    cleaned_text = text.strip()

    if len(cleaned_text) <= limit:
        return cleaned_text

    return cleaned_text[:limit] + "..."
